# -*- coding: utf-8 -*-
"""RUH 出货手续包 v3 修复版——全套对齐 176-61334210 参考版（数字式手机变体）。
修复点：
  1. 托运人声明：剥离 sdt 控件外壳为纯文本
  2. 应急措施：品名字号 7pt → 9pt
  3. 托书：填品名（英文 5pt + 中文 4.8pt，参考版格式）
  4. 运单副本：英文改用 5G DIGITAL MOBILE 变体 + 加中文品名段
  5. 安检单/应急措施/交运单品名：统一改为「数字式手机/5G数字移动电话机」变体
  6. 安检单品名仍 5.2pt（受品名栏 ~110pt 宽物理限制）
v4（2026-08-19 本机适配，PyMuPDF 1.28）：
  7. mixed_insert 改用 Font.text_length 推进 x（insert_text 新版返回 1 而非宽度）
  8. 安检单品名改 TextWriter + SimSun 字体对象（宋体半宽西文，对齐参考版 AdobeSongStd；
     内置 china-s 西文全宽超格；insert_text(fontfile=) 丢 ToUnicode 提取变点）
  9. 安检单补杂项品名行；split_goods 折行止于「电池」关键字
 10. 托书英文段起点基线 362.3 / 中文段 440.3（旧 357.2/435.4 为块顶，压标题）
"""
import fitz, os, re, shutil, zipfile
import openpyxl

_HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(_HERE, '..', '..', 'templates'))  # 相对脚本定位，跨平台可部署
PUB = os.path.join(ROOT, '公共模板')
TMPL2 = os.path.join(ROOT, '测试2')
JY_TPL = os.path.join(ROOT, '测试176-61334210模板')
OUT = os.path.join(os.path.dirname(ROOT), '出货手续包')

# 安检单品名用宋体（对齐参考版 AdobeSongStd、半宽西文）。Windows 用系统宋体；
# Linux 服务器把 simsun.ttc 放 templates/fonts/，或装思源宋体/文泉驿（按序找到即用）。
_FONT_CANDS = [
    r'C:/Windows/Fonts/simsun.ttc',
    os.path.join(ROOT, 'fonts', 'simsun.ttc'),
    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
    '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
]
_song_path = next((p for p in _FONT_CANDS if os.path.isfile(p)), None)


def song_font():
    if _song_path:
        return fitz.Font(fontfile=_song_path)
    print('警告：未找到宋体字体文件，安检单退回内置 china-s（西文全宽可能超格）；'
          'Linux 请安装思源宋体或把 simsun.ttc 放 templates/fonts/')
    return fitz.Font('china-s')

MASTERS = {
    '176-61333915': {
        'models': ['A2636', 'A2881', 'A2889', 'A3089', 'A3101', 'A3286', 'A3295', 'A3408'],
        'battery': ['A2656', 'A2863', 'A2866', 'A3018', 'A3011', 'NA264', '102CS', 'NA248H'],
        'pcs': 7,
        'hawbs': [('TYN00162000', 1, 42.0), ('TYN00161978', 3, 34.4),
                  ('TYN00161970', 1, 10.0), ('TYN00162026', 1, 16.4),
                  ('TYN00161979', 1, 8.0)],
    },
    '176-61333926': {
        'models': ['A2403', 'A2407', 'A3089', 'A3101', 'A3286', 'A3295', 'A3408'],
        'battery': ['A2479', 'A3018', 'A3011', 'NA264', '102CS', 'NA248H'],
        'pcs': 4,
        'hawbs': [('TYN00162013', 1, 17.2), ('TYN00161992', 3, 78.4)],
    },
}
DEST = 'RUH'
FLIGHT = 'EK309'      # 航班（独立版由邮件提取覆盖）
CARRIER = 'EK'        # 承运人（航班前两字母）
DATE_STR = '2026-08-16'  # 运输时间（独立版由邮件提取覆盖）
EN_HEAD = 'CONSOLIDATION AS PER ATTACHED MANIFEST'
EN_TAIL = 'LITHIUM ION BATTERIES IN COMPLIANCE WITH SECTION II OF PI 967,'

# 数字式手机/5G数字移动电话机变体（参考版 176-61334210 用此）
CN_GOODS = {
    'A2403': '数字式手机 A2403（内置聚合物锂离子电池 A2479, 3.83V 2815mAh 10.78Wh）',
    'A2407': '数字式手机 A2407 (内置聚合物锂离子电池 A2479, 3.83V 2815mAh 10.78Wh)',
    'A2636': '数字式手机 A2636（内置聚合物锂离子电池A2656,3.87V 3095mAh 11.97Wh）',
    'A2881': '数字式手机 A2881 (内置聚合物锂离子电池 APN: A2863 3.87V 3279mAh 12.68Wh)',
    'A2889': '数字式手机A2889 (内置聚合物锂离子电池 APN: A2866, 3.87V 3200mAh 12.38Wh)',
    'A3089': '5G数字移动电话机A3089 (内置聚合物锂离子电池 APN:A3018，3.876V 3349mAh 12.981Wh)',
    'A3101': '5G数字移动电话机A3101 (内置聚合物锂离子电池 APN:A3011，3.879V 3274mAh 12.700Wh)',
    'A3286': '5G数字移动电话机A3286 (内置聚合物锂离子电池 NA264, 3.886V, 13.839Wh, 3561mAh)',
    'A3295': '5G数字移动电话机A3295 (内置聚合物锂离子电池 102CS，3.878V 4685mAh 18.169Wh)',
    'A3408': '数字式手机 A3408（内置聚合物锂离子电池 NA248H, 3.884V 4005mAh 15.556Wh）',
}
EN_GOODS = {
    'A2403': 'DIGITAL CELLPHONE FOR RADIOTELEPHONY A2403 (BUILT WITH LITHIUM ION POLYMER BATTERY, A2479, 3.83V 2815mAh 10.78Wh)',
    'A2407': 'DIGITAL CELLPHONE FOR RADIOTELEPHONY A2407 (BUILT WITH LITHIUM ION POLYMER BATTERY A2479, 3.83V 2815mAh 10.78Wh)',
    'A2636': 'DIGITAL CELLPHONE FOR RADIOTELEPHONY A2636 (BUILT WITH LITHIUM ION POLYMER BATTERY, A2656,3.87V 3095mAh 11.97Wh)',
    'A2881': 'DIGITAL CELLPHONE FOR RADIOTELEPHONY A2881 (BUILT WITH LITHIUM ION POLYMER BATTERY, APN: A2863 3.87V 3279mAh 12.68Wh)',
    'A2889': 'DIGITAL CELLPHONE FOR RADIOTELEPHONY A2889（BUILT WITH LITHIUM ION POLYMER BATTERY, APN: A2866, 3.87V 3200mAh 12.38Wh）',
    'A3089': '5G DIGITAL MOBILE PHONE FOR RADIOTELEPHONY A3089（BUILT WITH LITHIUM ION POLYMER BATTERY，APN:A3018，3.876V 3349mAh 12.981Wh)',
    'A3101': '5G DIGITAL MOBILE PHONE FOR RADIOTELEPHONY A3101（BUILT WITH LITHIUM ION POLYMER BATTERY，APN:A3011, 3.879V 3274mAh 12.700Wh)',
    'A3286': '5G DIGITAL MOBILE PHONE FOR RADIOTELEPHONY A3286（BUILT WITH LITHIUM ION POLYMER BATTERY, NA264, 3.886V, 13.839Wh, 3561mAh)',
    'A3295': '5G DIGITAL MOBILE PHONE FOR RADIOTELEPHONY A3295（BUILT WITH LITHIUM ION POLYMER BATTERY，102CS，3.878V 4685mAh 18.169Wh)',
    'A3408': 'DIGITAL CELLPHONE FOR RADIOTELEPHONY A3408 (BUILT WITH LITHIUM ION POLYMER BATTERY NA248H, 3.884V 4005mAh 15.556Wh)',
}

# ================== 通用：中英混排 ==================
def _is_cjk(ch):
    o = ord(ch)
    return (0x2E80 <= o <= 0x9FFF or 0xF900 <= o <= 0xFAFF or
            0xFF00 <= o <= 0xFFEF or 0x3000 <= o <= 0x303F)

def mixed_insert(pg, x, y, text, size):
    segs = re.findall(r'[\u2e80-\u9fff\uf900-\ufaff\uff00-\uffef\u3000-\u303f]+|[^\u2e80-\u9fff\uf900-\ufaff\uff00-\uffef\u3000-\u303f]+', text)
    xp = x
    for seg in segs:
        if not seg: continue
        fn = 'china-s' if _is_cjk(seg[0]) else 'helv'
        pg.insert_text((xp, y), seg, fontname=fn, fontsize=size, color=0)
        # PyMuPDF>=1.24 左右 insert_text 返回 1 而非宽度，改用 text_length 推进 x（新旧版本通用）
        xp += fitz.Font(fn).text_length(seg, fontsize=size)
    return xp - x

def split_goods(s):
    kw = '内置聚合物锂离子电池'
    i = s.find(kw)
    line1 = s[:i + len(kw)]
    line2 = s[i + len(kw):].lstrip() + ','
    return line1, line2

def goods_of(info):
    """品名条目列表：优先 info['entries']（按邮件手机/模组分类挑好的 (中,英) 对），
    缺省回退按机型查内置 CN_GOODS/EN_GOODS。"""
    if info.get('entries'):
        return [c for c, e in info['entries']], [e for c, e in info['entries']]
    return ([CN_GOODS[m] for m in info['models']],
            [EN_GOODS[m] for m in info['models']])

def parts_of(info):
    """零件品名：一律以邮件非电池分单提取为准；邮件没有就不写，不做兜底。"""
    return info.get('parts') or []

def mixed_width(text, fs):
    """中英混排串宽：china-s/helv 分段量宽（与 mixed_insert 同口径）。"""
    tot = 0.0
    cur, cur_cjk = '', None
    for ch in text:
        isc = _is_cjk(ch)
        if cur_cjk is None or isc == cur_cjk:
            cur += ch
            cur_cjk = isc
        else:
            tot += fitz.Font('china-s' if cur_cjk else 'helv').text_length(cur, fontsize=fs)
            cur, cur_cjk = ch, isc
    if cur:
        tot += fitz.Font('china-s' if cur_cjk else 'helv').text_length(cur, fontsize=fs)
    return tot

# ================== 1. 托运人声明（剥 sdt 为纯文本） ==================
def make_shengming(master, info):
    src = os.path.join(PUB, '2026国际出港电池货物托运人声明.docx')
    z = zipfile.ZipFile(src)
    xml = z.read('word/document.xml').decode('utf-8')
    z.close()
    vals = [master, ', '.join(info['battery']), str(info['pcs'])]
    # 1) 替换占位符文本
    for v in vals:
        xml = xml.replace('<w:t>单击此处输入文字。</w:t>', '<w:t>%s</w:t>' % v, 1)
    # 2) 剥离 sdt 控件外壳为纯文本 run（消除 Word 里控件方框/括号渲染）
    # 每个 sdt 块形如 <w:sdt>...<w:sdtContent><w:permStart .../><w:r><w:t>VAL</w:t></w:r><w:permEnd.../></w:sdtContent></w:sdt>
    # 提取 sdtContent 里的 <w:r>...VAL...</w:r>，去掉 permStart/permEnd，替换 sdt 块
    pattern = re.compile(
        r'<w:sdt><w:sdtPr>.*?</w:sdtPr><w:sdtEndPr>.*?</w:sdtEndPr><w:sdtContent>'
        r'(<w:permStart[^/]*/>)?'
        r'(<w:r>.*?</w:r>)'
        r'(<w:permEnd[^/]*/>)?'
        r'</w:sdtContent></w:sdt>',
        re.S)
    def _strip(m):
        inner_r = m.group(2)
        # 去掉 permStart/permEnd（已在 .group(1)(3) 排除，但 inner 可能有）
        return inner_r
    xml_new = pattern.sub(_strip, xml)
    n_stripped = xml_new.count('<w:r>') != xml.count('<w:r>')  # 粗略
    out = os.path.join(OUT, master, master + ' 2026国际出港电池货物托运人声明.docx')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    stale = os.path.join(OUT, master, master + '  2026国际出港电池货物托运人声明.docx')  # NBSP 旧版
    if os.path.exists(stale):
        os.remove(stale)
    zin = zipfile.ZipFile(src)
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                data = xml_new.encode('utf-8')
            zout.writestr(item, data)
    zin.close()
    return out

# ================== 2. 安检单 ==================
def make_anxian(master, info):
    doc = fitz.open(os.path.join(PUB, '安检单.pdf'))
    pg = doc[0]
    mixed_insert(pg, 274.9, 186.5, master, 11)
    mixed_insert(pg, 96.4, 211.0, DEST, 11)
    mixed_insert(pg, 274.9, 211.0, CARRIER, 11)
    mixed_insert(pg, 96.4, 235.9, FLIGHT, 11)
    mixed_insert(pg, 275.7, 235.9, DATE_STR, 11)
    # 品名条目（按邮件手机/模组分类挑好的变体）——TextWriter + SimSun 整行写入。
    # 参考版用 AdobeSongStd-Light（宋体、半宽西文）；内置 china-s 西文全宽超格；
    # insert_text(fontfile=) 嵌 CID 字体丢 ToUnicode（提取变点），TextWriter 则文本层完整。
    cn_list, _ = goods_of(info)
    song = song_font()
    tw = fitz.TextWriter(pg.rect)
    parts = parts_of(info)
    MISC = (','.join(c for c, e in parts) + ',') if parts else ''   # 零件品名：只写邮件提取的，没有就不写
    W = 204.5 - 90.7

    def misc_lines(ff):
        ls, cur = [], ''
        for ch in MISC:
            if song.text_length(cur + ch, fontsize=ff) > W:
                ls.append(cur); cur = ch
            else:
                cur += ch
        if cur:
            ls.append(cur)
        return ls

    # 字号双向自适应：内容少→放大(≤6.5)；标准格(到357)放不下→擦357横线借空格(到406)；再不下→缩(≥4.5)
    CAP1, CAP2 = 357 - 243.5, 405 - 243.5
    pairs = [split_goods(s) for s in cn_list]

    def layout(f, cap):
        lh = f * 6.3 / 5.2
        ml = misc_lines(f)
        N = len(pairs) * 2 + len(ml)
        if N * lh > cap:
            return None
        wmax = max([song.text_length(t, fontsize=f) for t in ml] +
                   [song.text_length(t, fontsize=f) for p in pairs for t in p], default=0)
        if wmax > W:
            return None
        return lh, ml

    lay = None
    extended = False
    fs = None
    for f in (6.5, 6.0, 5.6, 5.2):
        lay = layout(f, CAP1)
        if lay:
            fs = f
            break
    if lay is None:
        extended = True
        for f in (5.6, 5.2, 4.9, 4.5):
            lay = layout(f, CAP2)
            if lay:
                fs = f
                break
    if lay is None:
        fs = 4.5
        ml = misc_lines(fs)
        N = len(pairs) * 2 + len(ml)
        lay = (CAP2 / N, ml)
    if extended:
        # 擦品名格底横线段（x91~203，y≈357），与下方空格合并；左右竖线本就通到 406
        pg.draw_rect(fitz.Rect(91, 355.5, 203, 358.5), color=None, fill=(1, 1, 1))
    lh, ml = lay
    y = 243.5 + fs * 0.8
    for l1, l2 in pairs:
        tw.append((90.7, y), l1, font=song, fontsize=fs)
        tw.append((90.7, y + lh), l2, font=song, fontsize=fs)
        y += 2 * lh
    for ln in ml:
        tw.append((90.7, y), ln, font=song, fontsize=fs)
        y += lh
    tw.write_text(pg, color=(0, 0, 0))
    out = os.path.join(OUT, master, master + ' 安检单.pdf')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.subset_fonts()  # 嵌入字体只留用到的字符子集（否则 simsun/china-s 全量嵌入 3~17MB）
    doc.save(out, garbage=3, deflate=True)
    doc.close()
    return out

# ================== 3. 应急措施（品名 9pt） ==================
def make_yingji(master, info):
    doc = fitz.open(os.path.join(PUB, '应急措施.pdf'))
    pg = doc[0]
    mixed_insert(pg, 77.0, 92.0, master, 12.6)
    # 品名每行一个条目；字号双向自适应：放得下放大(≤8)，放不下缩(≥6)；宽≤305、高≤126
    cn_list, _ = goods_of(info)
    CAP, W = 225 - 98.8, 505 - 199.8
    fs = lh = None
    for f in (8.0, 7.5, 7.0):
        l = f * 8.4 / 7.0
        if len(cn_list) * l <= CAP and max((mixed_width(c + ',', f) for c in cn_list), default=0) <= W:
            fs, lh = f, l
            break
    if fs is None:
        w7 = max((mixed_width(c + ',', 7.0) for c in cn_list), default=W) or W
        lh = min(8.4, CAP / max(len(cn_list), 1))
        fs = max(6.0, min(7.0, lh * 7.0 / 8.4, 7.0 * W / w7))
        lh = fs * 8.4 / 7.0
    y = 98.8 + fs * 0.8
    for cn in cn_list:
        mixed_insert(pg, 199.8, y, cn + ',', fs)
        y += lh
    out = os.path.join(OUT, master, master + ' 应急措施.pdf')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.subset_fonts()  # 嵌入字体只留用到的字符子集（否则 simsun/china-s 全量嵌入 3~17MB）
    doc.save(out, garbage=3, deflate=True)
    doc.close()
    return out

# ================== 4. 托书（填品名） ==================
def fill_tuoshū(master, info, consolidation=True):
    """托书品名：每个机型单独一行（参考用户提供的实例格式）
    consolidation=True 时英文段顶部加 'CONSOLIDATION AS PER ATTACHED MANIFEST'（用户要求默认有）。"""
    sdir = os.path.join(TMPL2, master + '模板')
    src = next(os.path.join(sdir, f) for f in os.listdir(sdir) if '托书' in f)
    doc = fitz.open(src)
    pg = doc[0]
    
    cn_list, en_list = goods_of(info)
    parts = parts_of(info)
    bpcs = info.get('bpcs', 0)
    
    # 构建英文段：每行一个机型
    en_lines = []
    if consolidation:
        en_lines.append(EN_HEAD + ',')
    for en in en_list:
        en_lines.append(en + ',')
    # 电池声明
    if bpcs:
        en_lines.append('LITHIUM ION BATTERIES IN COMPLIANCE WITH SECTION II OF PI 967,')
        en_parts = [e for c, e in parts if e]
        if en_parts:
            en_lines.append(','.join(en_parts) + ' LITHIUM BATTERY %d PCS' % bpcs)
        else:
            en_lines.append('LITHIUM BATTERY %d PCS' % bpcs)
    
    # 构建中文段：每行一个机型
    cn_lines = []
    for cn in cn_list:
        cn_lines.append(cn + ',')
    # 零件
    if parts:
        cn_parts = [c for c, e in parts if c]
        if cn_parts:
            cn_lines.append(','.join(cn_parts))
    
    # 排版参数
    helv = fitz.Font('helv')
    cjkf = fitz.Font('china-s')
    maxw = 575 - 251.7  # 可用宽度
    BOX_TOP, BOX_BOT = 362.3, 472.0
    
    # 字号：英文 5pt，中文 4.8pt，行距 6pt
    fs_e, lh_e = 5.0, 6.0
    fs_c, lh_c = 4.8, 5.8
    
    # 如果总高度超盒，按比例缩放（下限 0.8 倍）
    total_h = len(en_lines) * lh_e + len(cn_lines) * lh_c
    if total_h > BOX_BOT - BOX_TOP:
        sc = max(0.8, (BOX_BOT - BOX_TOP) / total_h)
        fs_e, lh_e = fs_e * sc, lh_e * sc
        fs_c, lh_c = fs_c * sc, lh_c * sc
    
    # 写入英文段
    y = BOX_TOP
    for line in en_lines:
        mixed_insert(pg, 251.7, y, line, fs_e)
        y += lh_e
    
    # 写入中文段
    for line in cn_lines:
        mixed_insert(pg, 251.7, y, line, fs_c)
        y += lh_c
    
    out = os.path.join(OUT, master, master + '托书.pdf')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.subset_fonts()
    doc.save(out, garbage=3, deflate=True)
    doc.close()
    return out

def wrap_mixed(text, default_font, fs, maxw):
    """按脚本切段 wrap，返回行列表。使用 china-s/helv 分别量宽。"""
    cjk = fitz.Font('china-s'); lat = fitz.Font('helv')
    segs = re.findall(r'[\u2e80-\u9fff\uf900-\ufaff\uff00-\uffef\u3000-\u303f]+|[^\u2e80-\u9fff\uf900-\ufaff\uff00-\uffef\u3000-\u303f]+', text)
    # 拼成"段"：连续同字体段
    units = []
    cur = ''
    cur_cjk = None
    for s in segs:
        isc = _is_cjk(s[0])
        if cur_cjk is None or isc == cur_cjk:
            cur += s
            cur_cjk = isc
        else:
            units.append((cur_cjk, cur))
            cur = s
            cur_cjk = isc
    if cur:
        units.append((cur_cjk, cur))
    # 行 build
    lines, line = [], ''
    for isc, u in units:
        f = cjk if isc else lat
        w_u = f.text_length(u, fontsize=fs)
        cur_w = (cjk if any(_is_cjk(c) for c in line) else lat).text_length(line, fontsize=fs) if line else 0
        if cur_w + w_u > maxw:
            if line.strip():
                lines.append(line)
            # u 可能本身超 maxw，强制截断
            if w_u > maxw:
                cut = ''
                for ch in u:
                    f2 = cjk if _is_cjk(ch) else lat
                    if f2.text_length(cut + ch, fontsize=fs) > maxw:
                        lines.append(cut)
                        cut = ch
                    else:
                        cut += ch
                line = cut
            else:
                line = u
        else:
            line += u
    if line.strip():
        lines.append(line)
    return lines

# ================== 5. 运单副本（英文改参考版变体 + 加中文品名段） ==================
EN_GOODS_RECT = (407.7, 482.9, 578.0, 590.7)
CN_GOODS_RECT = (407.7, 602.0, 578.0, 662.5)  # 中文品名段；底边止于 662.9 表格横线（再往下是 Shipper certifies 证明条款格，redact 越界会误删条款文字）
FS_WB, LH_WB = 4.6, 5.6

def fill_waybill(master, info):
    sdir = os.path.join(TMPL2, master + '模板')
    src = next(os.path.join(sdir, f) for f in os.listdir(sdir) if '运单副本' in f or '运单复本' in f)
    doc = fitz.open(src); pg = doc[0]
    # 擦除英文品名区 + 中文品名区
    pg.add_redact_annot(fitz.Rect(*EN_GOODS_RECT), fill=None)
    pg.add_redact_annot(fitz.Rect(*CN_GOODS_RECT), fill=None)
    pg.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE, graphics=fitz.PDF_REDACT_LINE_ART_NONE)
    cn_list, en_list = goods_of(info)
    parts = parts_of(info)
    # 英文段尾追加：LITHIUM BATTERY n PCS（n=含电池分单件数合计）+ 邮件提取的 SVC 英文零件名
    en_text = EN_HEAD + ',' + ','.join(en_list) + '.' + EN_TAIL
    bpcs = info.get('bpcs', 0)
    if bpcs:
        seg = 'LITHIUM BATTERY %d PCS' % bpcs
        en_parts = [e for c, e in parts if e]
        if en_parts:
            seg += ',' + ','.join(en_parts)
        en_text += seg

    def fit(lines_n, fs0, lh0, cap):
        """行数超格高时按比例缩字号/行距。"""
        if lines_n * lh0 <= cap:
            return fs0, lh0
        lh = cap / lines_n
        return min(fs0, lh * fs0 / lh0), lh

    # 英文品名（双向自适应：先试放大 5.0/4.8/4.6，放不下回退缩字）
    helv = fitz.Font('helv')
    en_cap = EN_GOODS_RECT[3] - EN_GOODS_RECT[1]
    en_w = EN_GOODS_RECT[2] - EN_GOODS_RECT[0]
    en_lines = None
    for f in (5.0, 4.8, 4.6):
        e = wrap_text(en_text, helv, f, en_w)
        if len(e) * f * LH_WB / FS_WB <= en_cap:
            en_lines, fs_e, lh_e = e, f, f * LH_WB / FS_WB
            break
    if en_lines is None:
        en_lines = wrap_text(en_text, helv, FS_WB, en_w)
        fs_e, lh_e = fit(len(en_lines), FS_WB, LH_WB, en_cap)
        en_lines = wrap_text(en_text, helv, fs_e, en_w)
    y = EN_GOODS_RECT[1] + fs_e * 0.86
    tw = fitz.TextWriter(pg.rect)
    for ln in en_lines:
        tw.append((EN_GOODS_RECT[0], y), ln, font=helv, fontsize=fs_e)
        y += lh_e
    tw.write_text(pg, color=(0, 0, 0))
    # 中文品名（按分类变体 + 邮件提取的零件行；双向自适应字号）
    cn_text = ','.join(cn_list) + (',' + ','.join(c for c, e in parts) + ',' if parts else '')
    cjk = fitz.Font('china-s')
    cn_cap = CN_GOODS_RECT[3] - CN_GOODS_RECT[1]
    cn_w = CN_GOODS_RECT[2] - CN_GOODS_RECT[0]
    cn_lines = None
    for f in (5.0, 4.8, 4.6):
        c = wrap_text_cn(cn_text, cjk, f, cn_w)
        if len(c) * f * LH_WB / FS_WB <= cn_cap:
            cn_lines, fs_c, lh_c = c, f, f * LH_WB / FS_WB
            break
    if cn_lines is None:
        cn_lines = wrap_text_cn(cn_text, cjk, FS_WB, cn_w)
        fs_c, lh_c = fit(len(cn_lines), FS_WB, LH_WB, cn_cap)
        cn_lines = wrap_text_cn(cn_text, cjk, fs_c, cn_w)
    y = CN_GOODS_RECT[1] + fs_c * 0.86
    tw2 = fitz.TextWriter(pg.rect)
    for ln in cn_lines:
        tw2.append((CN_GOODS_RECT[0], y), ln, font=cjk, fontsize=fs_c)
        y += lh_c
    tw2.write_text(pg, color=(0, 0, 0))
    out = os.path.join(OUT, master, master + '运单副本.pdf')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    stale = os.path.join(OUT, master, master + ' 运单副本.pdf')
    if os.path.exists(stale):
        os.remove(stale)
    doc.subset_fonts()  # 嵌入字体只留用到的字符子集（否则 simsun/china-s 全量嵌入 3~17MB）
    doc.save(out, garbage=3, deflate=True)
    doc.close()
    return out

def wrap_text(text, font, fs, maxw):
    units = re.findall(r'[A-Za-z0-9][A-Za-z0-9.,:()+/\-]*|\s|.', text)
    lines, cur = [], ''
    for u in units:
        t = cur + u
        if font.text_length(t, fontsize=fs) > maxw:
            if cur.strip(): lines.append(cur.rstrip())
            cur = '' if u == ' ' else u
            while font.text_length(cur, fontsize=fs) > maxw and len(cur) > 1:
                k = len(cur) - 1
                while k > 1 and font.text_length(cur[:k], fontsize=fs) > maxw:
                    k -= 1
                lines.append(cur[:k])
                cur = cur[k:]
        else:
            cur = t
    if cur.strip(): lines.append(cur.rstrip())
    return lines

def wrap_text_cn(text, font, fs, maxw):
    """中文文本 wrap（按字符截断）"""
    lines, cur = [], ''
    for ch in text:
        t = cur + ch
        if font.text_length(t, fontsize=fs) > maxw:
            if cur: lines.append(cur)
            cur = ch
        else:
            cur = t
    if cur: lines.append(cur)
    return lines

# ================== 6. 交运单 ==================
def fill_jiaoyun(master, info):
    src = next(os.path.join(JY_TPL, f) for f in os.listdir(JY_TPL) if '交运单' in f)
    wb = openpyxl.load_workbook(src); ws = wb['表样']
    # 解除与数据区(第5行起)相交的合并单元格：合并区内非左上角格只读，写入会报
    # 'MergedCell' object attribute 'value' is read-only（如模板 B10:B11/B12:B13）
    for rng in [r for r in list(ws.merged_cells.ranges) if r.min_row >= 5]:
        ws.unmerge_cells(str(rng))
    pcs = sum(h[1] for h in info['hawbs'])
    wt = round(sum(h[2] for h in info['hawbs']), 1)
    goods_text = '\n'.join(goods_of(info)[0])
    mr = None
    for row in ws.iter_rows():
        for c in row:
            if c.column == 2 and c.value and re.fullmatch(r'\d{3}-\d{8}', str(c.value).strip()):
                mr = c.row
    assert mr
    ws.cell(row=mr, column=2, value=master)
    ws.cell(row=mr, column=3, value=DEST)
    ws.cell(row=mr, column=4, value=pcs)
    ws.cell(row=mr, column=5, value=wt)
    fc = ws.cell(row=mr, column=6, value=goods_text)
    fc.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical='top')
    ws.cell(row=mr, column=7, value=f'{len(info["hawbs"])}分')
    n = len(info['hawbs'])
    for i, (hawb, hp, hw) in enumerate(info['hawbs']):
        r = 5 + i
        ws.cell(row=r, column=2, value=hawb)
        ws.cell(row=r, column=3, value=DEST)
        ws.cell(row=r, column=4, value=hp)
        ws.cell(row=r, column=5, value=hw)
    for r in range(5 + n, 10):
        for col in (2, 3, 4, 5):
            ws.cell(row=r, column=col).value = None
    out = os.path.join(OUT, master, f'交运单_{master}.xlsx')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    wb.save(out)
    return out

# ================== 7. 危险品确认单（docx 模板填充） ==================
WEIXIAN_ORG = '信诺'  # 鉴定机构：鉴定汇总无此字段，用常量，可改

def make_weixian(master, info):
    """危险品同意承运证明：填货单号/申报品名/鉴定报告编号/鉴定物品名称/鉴定机构。
    申报品名＝鉴定物品名称＝中文品名条目（，分隔）；
    报告编号取 info['cert_codes']（run_package 按条目顺序注入，与品名一一对应）。"""
    import docx
    import copy
    cn_list, _ = goods_of(info)
    goods = '，'.join(cn_list)
    codes = info.get('cert_codes') or []
    codes_str = '，'.join(map(str, codes))
    out = os.path.join(OUT, master, master + ' 危险品确认单.docx')
    os.makedirs(os.path.dirname(out), exist_ok=True)

    def set_runs(p, segs):
        """segs=[(text, 是否下划线),...]：重写段落为多个 run，沿用首 run 字体样式，
        填充内容带下划线（写在表单横线上的效果）。"""
        if not p.runs:
            p.add_run()
        first = p.runs[0]
        for r in p.runs[1:]:
            r._r.getparent().remove(r._r)
        first.text = segs[0][0]
        first.underline = segs[0][1]
        rpr = first._r.rPr
        for text, udl in segs[1:]:
            r = p.add_run(text)
            if rpr is not None:
                r._r.insert(0, copy.deepcopy(rpr))
            r.underline = udl

    def build(fs):
        """按字号 fs 生成：填充 + 排版（标题 18pt 加粗居中、正文 fs、1.25 倍行距、边距 72pt）。"""
        doc = docx.Document(os.path.join(PUB, '危险品确认单模板.docx'))
        for p in doc.paragraphs:
            t = p.text
            if '货单号：' in t:
                set_runs(p, [(re.sub(r'货单号：\s*，', '货单号： ', t), False),
                             (master, True), (' ，', False)])
            elif t.startswith('货物申报品名'):
                set_runs(p, [('货物申报品名：', False), (goods, True)])
            elif t.startswith('鉴定报告编号'):
                if codes:
                    set_runs(p, [('鉴定报告编号：', False), (codes_str, True)])
            elif t.startswith('鉴定物品名称'):
                set_runs(p, [('鉴定物品名称：', False), (goods, True)])
            elif t.startswith('鉴定机构'):
                set_runs(p, [('鉴定机构：   ', False), (WEIXIAN_ORG, True)])
        sec = doc.sections[0]
        sec.left_margin = sec.right_margin = docx.shared.Pt(72)
        for i, p in enumerate(doc.paragraphs):
            pf = p.paragraph_format
            if i == 0:
                if p.runs:
                    p.runs[0].text = p.runs[0].text.lstrip()  # 模板靠行首空格定位，不去掉会居右偏
                pf.alignment = 1
                pf.first_line_indent = docx.shared.Pt(0)  # 模板带首行缩进 7 字符，不清掉居中会偏右
                pf.left_indent = docx.shared.Pt(0)
                _ind = p._p.pPr.ind  # python-docx 清不掉 firstLineChars（字符缩进优先），直接删属性
                if _ind is not None:
                    _ind.attrib.pop('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLineChars', None)
                pf.space_after = docx.shared.Pt(10)
                for r in p.runs:
                    r.font.size = docx.shared.Pt(18)
                    r.bold = True
            elif i == 1:
                continue
            elif p.text.lstrip().startswith(('确认人', '单位（或公章）', '年')):
                for r in p.runs:
                    r.font.size = docx.shared.Pt(fs)
                pf.space_before = docx.shared.Pt(8)
            else:
                pf.line_spacing = 1.25
                pf.space_after = docx.shared.Pt(0)
                for r in p.runs:
                    r.font.size = docx.shared.Pt(fs)
        doc.save(out)

    # 一页内保守估算：行高按 1.6×fs（Word 中英混排实际行高），西文宽取 0.6fs
    AVAIL_W, AVAIL_H = 451.0, 698.0
    tpl_texts = [p.text for p in docx.Document(os.path.join(PUB, '危险品确认单模板.docx')).paragraphs]
    texts = []
    for t in tpl_texts:
        if '货单号：' in t:
            texts.append(re.sub(r'货单号：\s*，', '货单号： ', t) + master + ' ，')
        elif t.startswith('货物申报品名'):
            texts.append('货物申报品名：' + goods)
        elif t.startswith('鉴定报告编号'):
            texts.append('鉴定报告编号：' + codes_str)
        elif t.startswith('鉴定物品名称'):
            texts.append('鉴定物品名称：' + goods)
        elif t.lstrip().startswith(('确认人', '单位（或公章）', '年')) or not t.strip():
            continue
        else:
            texts.append(t)

    def est_h(fs):
        lh = fs * 1.55
        h = 18 * 1.4 + 10  # 标题行高+段后距
        for t in texts:
            w = sum(fs if '\u4e00' <= ch <= '\u9fff' or ch in '，：（）、《》' else fs * 0.6 for ch in t)
            h += max(1, -(-int(w) // int(AVAIL_W))) * lh
        h += 3 * (lh + 8)  # 落款三段
        return h

    def word_pages():
        """用 Word 实测页数（估算有偏差，以 Word 为准）；无 Word 返回 None。"""
        try:
            import win32com.client
            w = win32com.client.Dispatch('Word.Application')
            w.Visible = False
            w.DisplayAlerts = 0
            d = w.Documents.Open(out, ReadOnly=True)
            n = d.ComputeStatistics(2)
            d.Close(False)
            try:
                w.Quit()
            except Exception:
                pass
            return n
        except Exception:
            return None

    cands = [f for f in (14, 13, 12.5, 12, 11.5, 11, 10.5, 10) if est_h(f) <= AVAIL_H] or [10]
    for fs in cands:
        build(fs)
        n = word_pages()
        if n is None or n <= 1:  # 无 Word 时信保守估算；实测一页即停
            break
    return out

if __name__ == '__main__':
    for master, info in MASTERS.items():
        print('====', master, '====')
        print('  托运人声明:', os.path.basename(make_shengming(master, info)))
        print('  安检单:', os.path.basename(make_anxian(master, info)))
        print('  应急措施:', os.path.basename(make_yingji(master, info)))
        print('  托书:', os.path.basename(fill_tuoshū(master, info)))
        print('  运单副本:', os.path.basename(fill_waybill(master, info)))
        print('  交运单:', os.path.basename(fill_jiaoyun(master, info)))
        print('  危险品确认单:', os.path.basename(make_weixian(master, info)))
    print('ALL DONE')
